# coding:utf8

import json
import time

from docx import Document
from docx.shared import Inches

from monitor.test import matplotlib_test


class DocUtils(object):
    """
    DOC 数据解析
    """
    def __init__(self):
        pass

    def parseDataProduceDoc(self, data=None):
        """
        解析数据，并生成 DOC 格式的文件
        :param data:
        :return:
        """
        if data is None:
            return None

        data = json.loads(data)
        document = Document()
        # title
        document.add_heading(u'巡检日志', 0)

        ctime = time.time()
        for server_ip in data:
            document.add_heading(u'Service IP：' + server_ip, 1)
            for server_type in data[server_ip]:
                for server_name in data[server_ip][server_type]:
                    document.add_heading(u'%s' % server_name, 1)
                    serverData = data[server_ip][server_type][server_name]
                    # 字符串中指定的服务，使用柱状图或线形图
                    if server_name.lower() in 'mem,':
                        plats = matplotlib_test.MatplotlibUtils()
                        plats.parse(serverData, '%s.png' % ctime)
                    else:
                        # 判断类型，DISK的值是list，不是dict
                        if isinstance(serverData, list):
                            self.tableRows(serverData, document)
                        else:
                            self.table(serverData, document)
        document.add_picture('%s.png' % ctime, width=Inches(1.25))
        document.save(u'x_巡检日志.docx')

    def table(self, data, document):
        """
        创建表格,并写入一行数据
        :param data: 要写入的数据，类型 Dict
        :param document: 文档句柄
        :return:
        """
        # , style='Table Grid'
        table = document.add_table(rows=1, cols=len(data))
        table.autofit = False
        # 添加列
        index = 0
        tb_cells = table.rows[0].cells
        for server_index in data:
            tb_cells[index].text = server_index
            index += 1

        index = 0
        # 添加行
        row_cells = table.add_row().cells
        for server_index in data:
            server_data = data[server_index]
            if isinstance(server_data, list):
                self.tableRows(server_data, row_cells[index])
            else:
                row_cells[index].text = str(data[server_index])
            index += 1

    def tableRows(self, list_data, document):
        """
        创建表格,并写入多行数据
        :param list_data: 要写入的数据，类型 List
        :param document: 文档句柄
        :return:
        """
        table = document.add_table(rows=1, cols=len(list_data[0]))
        table.autofit = False

        # 添加列
        index = 0
        tb_cells = table.rows[0].cells
        for server_index in list_data[0]:
            tb_cells[index].text = server_index
            index += 1

        # 添加行
        row_cells = table.add_row().cells
        for data in list_data:
            index = 0
            for server_index in data:
                server_data = data[server_index]
                if isinstance(server_data, list):
                    self.tableRows(server_data, row_cells[index])
                else:
                    row_cells[index].text = str(data[server_index])
                index += 1


if __name__ == "__main__":
    doc = DocUtils()
    data = """
    {"127.0.0.1":s{"Service": {"Ufa": {"unode": "started", "CENTER_NODE_ADDR": "http://127.0.0.1:9090/uctr/stor",
    "uctr": "started"}, "Xserver": {"xserver.mysql.host": "jdbc:mysql://127.0.0.1", "server": "started"}, "Mysql":
    {"server": "started"}}, "Server": {"MEM": {"swap_used": "0B", "swap_free": "2.0G", "mem_buffers": "106.7M",
    "mem_used": "1.2G", "mem_total": "1.8G", "mem_used_rate": "65.17", "mem_free_rate": "34.83", "mem_cached":
    "505.7M", "swap_total": "2.0G", "mem_free": "652.6M", "swap_used_rate": "0.00", "swap_free_rate": "100.00"},
    "DISK": [{"used": "6.4G", "free": "10.9G", "fstype": "ext4", "dev": "/dev/mapper/VolGroup-lv_root", "path": "/",
    "total": "17.3G", "used_rate": "37.11"}, {"used": "38.7M", "free": "445.5M", "fstype": "ext4", "dev": "/dev/sda1",
     "path": "/boot", "total": "484.2M", "used_rate": "8.00"}], "CPU": {"cores": [{"model": "Intel(R) Core(TM)
     i7-6700HQ CPU @ 2.60GHz", "bits": "64bit"}], "cpu_count": 1, "core_count": 1}}}}
    """
    doc.parseDataProduceDoc(data)
