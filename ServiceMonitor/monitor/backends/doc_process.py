# coding:utf8

import json
import time

from docx import Document
from docx.shared import Inches

from monitor.backends.matplot_process import MatplotlibUtils


class DocUtils(object):
    """
    DOC 数据解析
    """
    def __init__(self):
        pass

    def parseDataProduceDoc(self, data=None):
        """
        解析数据，并生成 DOC 格式的文件
        :param data:
        :return:
        """
        if data is None:
            return None

        data = json.loads(data)

        document = Document()
        # title
        document.add_heading(u'巡检日志', 0)
        ctime = time.time()
        for server_ip in data:
            document.add_heading(u'Service IP：' + server_ip, 1)
            for server_type in data[server_ip]:
                for server_name in data[server_ip][server_type]:
                    document.add_heading(u'%s' % server_name, 1)
                    serverData = data[server_ip][server_type][server_name]
                    # 字符串中指定的服务，使用柱状图或线形图
                    if server_name.lower() in 'mem,':
                        plats = MatplotlibUtils()
                        plats.parse(serverData, '/tmp/%s.png' % ctime)
                        document.add_picture('/tmp/%s.png' % ctime, width=Inches(7.0))
                    else:
                        # 判断类型，DISK的值是list，不是dict
                        if isinstance(serverData, list):
                            self.tableRows(serverData, document)
                        else:
                            self.table(serverData, document)

        document.save(u'/tmp/x_巡检日志.docx')

    def table(self, data, document):
        """
        创建表格,并写入一行数据
        :param data: 要写入的数据，类型 Dict
        :param document: 文档句柄
        :return:
        """
        # , style='Table Grid'
        table = document.add_table(rows=1, cols=len(data))
        table.autofit = False
        # 添加列
        index = 0
        tb_cells = table.rows[0].cells
        for server_index in data:
            tb_cells[index].text = server_index
            index += 1

        index = 0
        # 添加行
        row_cells = table.add_row().cells
        for server_index in data:
            server_data = data[server_index]
            if isinstance(server_data, list):
                self.tableRows(server_data, row_cells[index])
            else:
                row_cells[index].text = str(data[server_index])
            index += 1

    def tableRows(self, list_data, document):
        """
        创建表格,并写入多行数据
        :param list_data: 要写入的数据，类型 List
        :param document: 文档句柄
        :return:
        """
        table = document.add_table(rows=1, cols=len(list_data[0]))
        table.autofit = False

        # 添加列
        index = 0
        tb_cells = table.rows[0].cells
        for server_index in list_data[0]:
            tb_cells[index].text = server_index
            index += 1

        # 添加行
        row_cells = table.add_row().cells
        for data in list_data:
            index = 0
            for server_index in data:
                server_data = data[server_index]
                if isinstance(server_data, list):
                    self.tableRows(server_data, row_cells[index])
                else:
                    row_cells[index].text = str(data[server_index])
                index += 1



